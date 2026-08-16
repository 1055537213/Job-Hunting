"""把证据约束简历草稿导出为可下载的 DOCX 和 PDF。

导出器不再调用模型，也不修改草稿内容。它只把已经通过真实性检查的正文应用到
一套紧凑、适合中文求职简历的排版上，并返回内存字节供文件版本服务持久化。
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from html import escape
from io import BytesIO

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from .resume_document import DOCX_MEDIA_TYPE, PDF_MEDIA_TYPE, sanitize_download_filename


@dataclass(frozen=True)
class GeneratedResumeFile:
    """一份尚未写入文件系统的导出文件。"""

    filename: str
    media_type: str
    content: bytes


# 基于 compact_reference_guide，并使用 A4/较窄页边距作为“中国求职简历”命名覆盖。
BODY_FONT_ASCII = "Calibri"
BODY_FONT_EAST_ASIA = "Microsoft YaHei"
HEADING_BLUE = RGBColor(0x2E, 0x74, 0xB5)
HEADING_DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED_GRAY = RGBColor(0x66, 0x66, 0x66)


def export_tailored_resume_files(
    *,
    candidate_name: str,
    job_title: str,
    draft_version: int,
    content: str,
) -> list[GeneratedResumeFile]:
    """用同一份安全正文生成 DOCX/PDF 两个可下载版本。"""

    stem = sanitize_download_filename(
        f"{candidate_name}_{job_title}_定制简历_v{draft_version}",
        fallback=f"tailored_resume_v{draft_version}",
    )
    return [
        GeneratedResumeFile(
            filename=f"{stem}.docx",
            media_type=DOCX_MEDIA_TYPE,
            content=build_docx(candidate_name, job_title, content),
        ),
        GeneratedResumeFile(
            filename=f"{stem}.pdf",
            media_type=PDF_MEDIA_TYPE,
            content=build_pdf(candidate_name, job_title, content),
        ),
    ]


def build_docx(candidate_name: str, job_title: str, content: str) -> bytes:
    """创建使用真实 Word 标题/列表样式的 A4 简历 DOCX。"""

    document = Document()
    configure_docx_page(document)
    configure_docx_styles(document)
    add_docx_title(document, candidate_name, job_title)

    for kind, text in parse_resume_blocks(content, candidate_name):
        if kind == "heading":
            paragraph = document.add_paragraph(text, style="Heading 1")
            paragraph.paragraph_format.keep_with_next = True
        elif kind == "bullet":
            paragraph = document.add_paragraph(text, style="List Bullet")
        else:
            document.add_paragraph(text, style="Normal")

    # 清理作者等默认元数据，避免下载文件携带开发机身份信息。
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.title = f"{candidate_name} - {job_title}"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def configure_docx_page(document: Document) -> None:
    """应用 A4 简历页面覆盖，保持稳定边距和可用宽度。"""

    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.7)
    section.right_margin = Inches(0.72)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.72)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)


def configure_docx_styles(document: Document) -> None:
    """把正文、标题和真实列表样式固定到简历排版 token。"""

    normal = document.styles["Normal"]
    set_style_font(normal, BODY_FONT_ASCII, BODY_FONT_EAST_ASIA, 10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading = document.styles["Heading 1"]
    set_style_font(heading, BODY_FONT_ASCII, BODY_FONT_EAST_ASIA, 13, bold=True)
    heading.font.color.rgb = HEADING_BLUE
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(5)
    heading.paragraph_format.line_spacing = 1.0

    bullet = document.styles["List Bullet"]
    set_style_font(bullet, BODY_FONT_ASCII, BODY_FONT_EAST_ASIA, 10.5)
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.25


def set_style_font(style, ascii_font: str, east_asia_font: str, size: float, bold: bool = False) -> None:
    """同时设置西文字体和东亚字体，避免 Word/LibreOffice 字体回退不一致。"""

    style.font.name = ascii_font
    style.font.size = Pt(size)
    style.font.bold = bold
    style._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)


def add_docx_title(document: Document, candidate_name: str, job_title: str) -> None:
    """添加克制的姓名与应聘方向标题块，不使用表格做布局。"""

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run(candidate_name)
    set_run_font(run, size=21, color=HEADING_DARK_BLUE, bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(9)
    run = subtitle.add_run(f"应聘方向：{job_title}")
    set_run_font(run, size=10.5, color=MUTED_GRAY)


def set_run_font(run, *, size: float, color: RGBColor, bold: bool = False) -> None:
    """给直接格式化的标题 run 同时设置中西文字体。"""

    run.font.name = BODY_FONT_ASCII
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT_ASCII)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT_ASCII)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT_EAST_ASIA)


def build_pdf(candidate_name: str, job_title: str, content: str) -> bytes:
    """使用 ReportLab 生成带中文字体的 A4 PDF。"""

    register_pdf_font()
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=17 * mm,
        bottomMargin=17 * mm,
        title=f"{candidate_name} - {job_title}",
        author="",
    )
    styles = pdf_styles()
    story = [
        Paragraph(escape(candidate_name), styles["ResumeTitle"]),
        Paragraph(escape(f"应聘方向：{job_title}"), styles["ResumeSubtitle"]),
        Spacer(1, 4),
    ]
    for kind, text in parse_resume_blocks(content, candidate_name):
        safe_text = escape(text).replace("\n", "<br/>")
        if kind == "heading":
            story.append(Paragraph(safe_text, styles["ResumeHeading"]))
        elif kind == "bullet":
            # CID 字体在不同 PDF 阅读器中可能丢失圆点符号，使用 ASCII 连字符更稳定。
            story.append(Paragraph(safe_text, styles["ResumeBullet"], bulletText="-"))
        else:
            story.append(Paragraph(safe_text, styles["ResumeBody"]))
    document.build(story)
    return output.getvalue()


def register_pdf_font() -> None:
    """注册 ReportLab 自带的简体中文 CID 字体，避免中文显示成方框。"""

    if "STSong-Light" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))


def pdf_styles() -> dict[str, ParagraphStyle]:
    """构造与 DOCX 相同层级和节奏的 PDF 样式表。"""

    base = getSampleStyleSheet()
    return {
        "ResumeTitle": ParagraphStyle(
            "ResumeTitle",
            parent=base["Normal"],
            fontName="STSong-Light",
            fontSize=21,
            leading=25,
            textColor=HexColor("#1F4D78"),
            spaceAfter=3,
            alignment=TA_LEFT,
        ),
        "ResumeSubtitle": ParagraphStyle(
            "ResumeSubtitle",
            parent=base["Normal"],
            fontName="STSong-Light",
            fontSize=10.5,
            leading=13,
            textColor=HexColor("#666666"),
            spaceAfter=9,
        ),
        "ResumeHeading": ParagraphStyle(
            "ResumeHeading",
            parent=base["Heading1"],
            fontName="STSong-Light",
            fontSize=13,
            leading=16,
            textColor=HexColor("#2E74B5"),
            spaceBefore=10,
            spaceAfter=5,
            keepWithNext=True,
        ),
        "ResumeBody": ParagraphStyle(
            "ResumeBody",
            parent=base["Normal"],
            fontName="STSong-Light",
            fontSize=10.5,
            leading=13.1,
            textColor=HexColor("#222222"),
            spaceAfter=6,
        ),
        "ResumeBullet": ParagraphStyle(
            "ResumeBullet",
            parent=base["Normal"],
            fontName="STSong-Light",
            fontSize=10.5,
            leading=13.1,
            textColor=HexColor("#222222"),
            leftIndent=14,
            firstLineIndent=0,
            bulletIndent=0,
            spaceAfter=4,
        ),
    }


def parse_resume_blocks(content: str, candidate_name: str) -> list[tuple[str, str]]:
    """把常见 Markdown/中文分节标记转换成导出器使用的轻量块。"""

    blocks: list[tuple[str, str]] = []
    for raw_line in content.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        heading_match = re.match(r"^#{1,6}\s+(.+)$", line)
        if heading_match:
            heading_text = heading_match.group(1).strip()
            # 模型经常把姓名作为第一个 H1；标题块已经展示姓名，因此不重复。
            if heading_text == candidate_name and not blocks:
                continue
            blocks.append(("heading", heading_text))
            continue
        if re.fullmatch(r"【[^】]+】", line):
            blocks.append(("heading", line[1:-1].strip()))
            continue
        bullet_match = re.match(r"^(?:[-*+]\s+|[•●]\s*)(.+)$", line)
        if bullet_match:
            blocks.append(("bullet", bullet_match.group(1).strip()))
            continue
        blocks.append(("paragraph", line))
    if not blocks:
        blocks.append(("paragraph", "暂无可导出的简历正文。"))
    return blocks
