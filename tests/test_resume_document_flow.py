"""简历文件上传、解析、改写和导出工作流测试。

这些测试把上传的二进制文件、PostgreSQL 元数据和 RAG 长文本登记视为同一条业务链路，
同时锁住账号隔离边界：知道别人的文件 ID 也不能列出或下载该文件。
"""

from __future__ import annotations

import zipfile
from io import BytesIO

import pytest
from docx import Document
from fastapi.testclient import TestClient
from PIL import Image
from reportlab.pdfgen import canvas

from job_hunting_agent import app as app_module
from job_hunting_agent import resume_document
from job_hunting_agent.app import JobHuntingApp
from job_hunting_agent.deduplication import DuplicateResourceError
from job_hunting_agent.llm import StaticLLMClient
from job_hunting_agent.models import CandidateProfileInput
from job_hunting_agent.resume_document import (
    ResumeDocumentError,
    ResumeExtraction,
    extract_resume_document,
    sanitize_download_filename,
)
from job_hunting_agent.web import create_web_app


def build_docx_bytes(*paragraphs: str) -> bytes:
    """在内存中创建测试 DOCX，避免把测试夹具写进仓库。"""

    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_text_pdf_bytes(text: str) -> bytes:
    """创建一个包含可直接提取 ASCII 文本的单页 PDF。"""

    output = BytesIO()
    pdf = canvas.Canvas(output)
    pdf.drawString(72, 760, text)
    pdf.save()
    return output.getvalue()


def build_scanned_pdf_bytes() -> bytes:
    """创建只有图片、没有 PDF 文本层的单页扫描件。"""

    image = Image.new("RGB", (900, 1200), "white")
    output = BytesIO()
    image.save(output, format="PDF")
    return output.getvalue()


def profile_input(name: str = "小林") -> CandidateProfileInput:
    """返回各测试共享的最小候选人结构化档案。"""

    return CandidateProfileInput(
        name=name,
        status="离职",
        education="本科",
        experience_years=1,
        skills={"Python": "项目使用"},
        preferred_cities=["杭州市"],
        salary_floor_k=10,
        expected_salary_k=15,
        target_directions=["Python 后端开发"],
        unacceptable=[],
    )


def register_and_login(client: TestClient, email: str) -> None:
    """为需要真实 Session Cookie 的 Web 测试创建并登录普通账号。"""

    password = "password-123"
    registered = client.post(
        "/api/auth/register",
        json={"email": email, "password": password, "display_name": email.split("@")[0]},
    )
    assert registered.status_code == 200
    logged_in = client.post("/api/auth/login", json={"email": email, "password": password})
    assert logged_in.status_code == 200


def create_profile_over_web(client: TestClient, name: str = "小林") -> int:
    """通过受认证 Web API 创建候选人，返回档案 ID。"""

    response = client.post(
        "/api/profiles",
        json={
            "name": name,
            "status": "离职",
            "education": "本科",
            "experience_years": 1,
            "skills": {"Python": "项目使用"},
            "preferred_cities": ["杭州市"],
            "salary_floor_k": 10,
            "expected_salary_k": 15,
            "target_directions": ["Python 后端开发"],
            "unacceptable": [],
        },
    )
    assert response.status_code == 200
    return int(response.json()["candidate_id"])


def test_resume_parser_handles_docx_text_pdf_and_scanned_pdf() -> None:
    """解析器应按真实文件内容区分 DOCX、文字 PDF 和需要 OCR 的扫描 PDF。"""

    docx_result = extract_resume_document(
        "candidate.docx",
        build_docx_bytes("小林", "Python 与 FastAPI 项目经历"),
    )
    text_pdf_result = extract_resume_document(
        "candidate.pdf",
        build_text_pdf_bytes("Python backend resume with FastAPI project experience"),
    )
    scanned_pdf_result = extract_resume_document(
        "scan.pdf",
        build_scanned_pdf_bytes(),
        ocr_runner=lambda _image: "扫描简历：Python、FastAPI 求职助手项目开发与接口设计经历",
    )

    assert docx_result.method == "docx"
    assert "FastAPI" in docx_result.text
    assert text_pdf_result.method == "pdf_text"
    assert "backend resume" in text_pdf_result.text
    assert scanned_pdf_result.method == "pdf_ocr"
    assert "扫描简历" in scanned_pdf_result.text
    assert scanned_pdf_result.page_count == 1


def test_pdf_text_layer_inspection_does_not_start_ocr(monkeypatch) -> None:
    """扫描 PDF 的上传前检查只能判断是否需要 OCR，不能在 Web 进程运行 OCR。"""

    def fail_if_ocr_runs(_image):
        """OCR 若在检查阶段运行，测试必须失败。"""

        raise AssertionError("PDF 文本层检查不应调用 RapidOCR")

    monkeypatch.setattr(resume_document, "run_rapidocr", fail_if_ocr_runs)

    inspection = resume_document.inspect_pdf_for_ocr(build_scanned_pdf_bytes())

    assert inspection.page_count == 1
    assert inspection.pages_needing_ocr == [0]


@pytest.mark.parametrize(
    ("filename", "content"),
    [
        ("resume.txt", b"plain text is not an accepted resume file"),
        ("resume.docx", b"not a zip document"),
        ("resume.pdf", b"not a pdf document"),
    ],
)
def test_resume_parser_rejects_unsupported_or_corrupt_files(filename: str, content: bytes) -> None:
    """扩展名不支持或文件签名损坏时必须拒绝，不能把任意文本当简历保存。"""

    with pytest.raises(ResumeDocumentError):
        extract_resume_document(filename, content)


def test_docx_parser_rejects_abnormally_large_uncompressed_xml(monkeypatch) -> None:
    """DOCX 在读取 XML 前必须限制解压大小，避免小压缩包异常膨胀。"""

    output = BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "<Types />")
        archive.writestr(
            "word/document.xml",
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            + ("A" * 2_000)
            + "</w:document>",
        )
    monkeypatch.setattr(resume_document, "MAX_DOCX_XML_PART_BYTES", 1_000)

    with pytest.raises(ResumeDocumentError, match="结构异常过大"):
        extract_resume_document("resume.docx", output.getvalue())


def test_sanitized_long_resume_filename_keeps_supported_extension() -> None:
    """超长下载名被截断后仍应保留扩展名，避免合法简历被误判格式。"""

    filename = sanitize_download_filename(("候选人" * 80) + ".docx")

    assert len(filename) <= 120
    assert filename.endswith(".docx")


def test_app_saves_uploaded_resume_as_versioned_artifact_and_rag_source(tmp_path) -> None:
    """上传简历应同时保存原文件、受控元数据和一条候选人范围的长文本来源。"""

    app = JobHuntingApp( resume_dir=tmp_path / "resume-files")
    app.initialize()
    account_a = app.auth.register("a@example.com", "password-123")
    account_b = app.auth.register("b@example.com", "password-123")
    candidate_id = app.save_candidate_profile(profile_input(), account_id=account_a.id)
    original = build_docx_bytes("小林", "Python 与 FastAPI 项目经历")

    first = app.upload_resume_document(
        candidate_id,
        "小林简历.docx",
        original,
        account_id=account_a.id,
    )
    second = app.upload_resume_document(
        candidate_id,
        "小林简历-更新.docx",
        build_docx_bytes("小林", "Python、FastAPI 与 LangChain 项目经历"),
        account_id=account_a.id,
    )

    assert first.artifact_type == "source"
    assert first.version == 1
    assert second.version == 2
    assert first.sha256 != second.sha256
    assert app.resume_file_path(first).read_bytes() == original
    assert [item.id for item in app.list_resume_artifacts(candidate_id, account_id=account_a.id)] == [
        first.id,
        second.id,
    ]
    indexed_sources = app.store.list_long_texts(
        entity_types=["resume_artifact"],
        account_id=account_a.id,
        candidate_id=candidate_id,
    )
    assert len(indexed_sources) == 2
    assert indexed_sources[0].entity_id == first.id
    assert "FastAPI" in indexed_sources[0].text

    with pytest.raises(KeyError):
        app.list_resume_artifacts(candidate_id, account_id=account_b.id)
    with pytest.raises(KeyError):
        app.get_resume_artifact(first.id, account_id=account_b.id)


def test_app_rejects_duplicate_resume_bytes_for_the_same_candidate(tmp_path) -> None:
    """同一候选人重复上传相同文件时，不应再创建新版本或额外文件。"""

    app = JobHuntingApp(resume_dir=tmp_path / "resume-files")
    app.initialize()
    account = app.auth.register("duplicate-resume@example.com", "password-123")
    candidate_id = app.save_candidate_profile(profile_input(), account_id=account.id)
    content = build_docx_bytes("小林", "Python 与 FastAPI 项目经历")

    first = app.upload_resume_document(
        candidate_id,
        "resume.docx",
        content,
        account_id=account.id,
    )
    with pytest.raises(DuplicateResourceError, match="简历"):
        app.upload_resume_document(
            candidate_id,
            "same-content-renamed.docx",
            content,
            account_id=account.id,
        )

    assert [item.id for item in app.list_resume_artifacts(candidate_id, account_id=account.id)] == [first.id]


def test_same_resume_bytes_can_belong_to_another_candidate_in_shared_account(tmp_path) -> None:
    """共享账号的不同候选人可以各自保存同一份源简历。"""

    app = JobHuntingApp(resume_dir=tmp_path / "resume-files")
    app.initialize()
    account = app.auth.register("shared-resume@example.com", "password-123")
    first_candidate_id = app.save_candidate_profile(profile_input("小林"), account_id=account.id)
    second_candidate_id = app.save_candidate_profile(profile_input("小周"), account_id=account.id)
    content = build_docx_bytes("共享模板", "Python 与 FastAPI 项目经历")

    first = app.upload_resume_document(
        first_candidate_id,
        "resume.docx",
        content,
        account_id=account.id,
    )
    second = app.upload_resume_document(
        second_candidate_id,
        "resume.docx",
        content,
        account_id=account.id,
    )

    assert first.id != second.id
    assert first.candidate_id == first_candidate_id
    assert second.candidate_id == second_candidate_id


def test_app_defers_scanned_pdf_ocr_then_registers_one_rag_source(tmp_path, monkeypatch) -> None:
    """待处理扫描 PDF 先保存原件，Worker 成功后才登记一次长文本来源。"""

    app = JobHuntingApp(resume_dir=tmp_path / "resume-files")
    app.initialize()
    account = app.auth.register("deferred-ocr@example.com", "password-123")
    candidate_id = app.save_candidate_profile(profile_input(), account_id=account.id)
    pending = app.upload_resume_document(
        candidate_id,
        "scan.pdf",
        build_scanned_pdf_bytes(),
        account_id=account.id,
        defer_ocr=True,
    )

    assert pending.status == "processing"
    assert pending.extraction_method == "pending_ocr"
    assert pending.long_text_id is None
    assert app.store.list_long_texts(
        entity_types=["resume_artifact"],
        account_id=account.id,
        candidate_id=candidate_id,
    ) == []

    monkeypatch.setattr(
        app_module,
        "extract_resume_document",
        lambda _filename, _content: ResumeExtraction(
            text="扫描简历 OCR 得到 Python、FastAPI 与 PostgreSQL 项目经验。",
            method="pdf_ocr",
            page_count=1,
        ),
    )
    completed = app.process_resume_ocr_artifact(
        artifact_id=pending.id,
        account_id=account.id,
        candidate_id=candidate_id,
    )
    repeated = app.process_resume_ocr_artifact(
        artifact_id=pending.id,
        account_id=account.id,
        candidate_id=candidate_id,
    )

    assert completed.status == "ready"
    assert completed.extraction_method == "pdf_ocr"
    assert completed.long_text_id is not None
    assert repeated.long_text_id == completed.long_text_id
    long_texts = app.store.list_long_texts(
        entity_types=["resume_artifact"],
        account_id=account.id,
        candidate_id=candidate_id,
    )
    assert len(long_texts) == 1
    assert "PostgreSQL" in long_texts[0].text


def test_web_upload_list_and_download_are_scoped_to_logged_in_account(tmp_path) -> None:
    """Web 下载接口必须再次校验账号，不能仅凭可枚举的 artifact_id 返回文件。"""

    web_app = create_web_app(
        resume_dir=tmp_path / "resume-files",
    )
    owner = TestClient(web_app)
    stranger = TestClient(web_app)
    register_and_login(owner, "owner@example.com")
    register_and_login(stranger, "stranger@example.com")
    candidate_id = create_profile_over_web(owner)
    original = build_docx_bytes("小林", "Python 与 FastAPI 项目经历")

    uploaded = owner.post(
        "/api/resumes/upload",
        data={"candidate_id": str(candidate_id)},
        files={
            "file": (
                "resume.docx",
                original,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert uploaded.status_code == 200
    artifact = uploaded.json()["artifact"]
    assert artifact["extraction_method"] == "docx"
    assert artifact["download_url"].endswith(f"/{artifact['id']}/download")
    assert "storage_key" not in artifact
    assert "account_id" not in artifact
    listed = owner.get("/api/resumes", params={"candidate_id": candidate_id})
    downloaded = owner.get(artifact["download_url"])
    forbidden = stranger.get(artifact["download_url"])

    assert [item["id"] for item in listed.json()["artifacts"]] == [artifact["id"]]
    assert downloaded.status_code == 200
    assert downloaded.content == original
    assert forbidden.status_code == 404

    stranger_delete = stranger.delete(f"/api/resumes/{artifact['id']}")
    deleted = owner.delete(f"/api/resumes/{artifact['id']}")
    listed_after_delete = owner.get("/api/resumes", params={"candidate_id": candidate_id})
    missing_download = owner.get(artifact["download_url"])

    assert stranger_delete.status_code == 404
    assert deleted.status_code == 200
    assert deleted.json()["deleted"] is True
    assert listed_after_delete.json()["artifacts"] == []
    assert missing_download.status_code == 404


def test_web_upload_enqueues_rag_index_when_worker_is_enabled(tmp_path, monkeypatch) -> None:
    """队列开启时上传接口立即返回任务，不在 Web 请求内调用 Embedding。"""

    class FakeQueue:
        """不连接 Redis 的最小队列替身，保留投递记录供断言。"""

        def __init__(self) -> None:
            self.enqueued: list[str] = []

        def health_check(self) -> None:
            """测试替身始终可用。"""

        def enqueue(self, task_key: str) -> None:
            """记录 Web 交给 Worker 的任务键。"""

            self.enqueued.append(task_key)

    def fail_if_called(*_args, **_kwargs):
        """如果 Web 同步索引，测试应立即失败。"""

        raise AssertionError("队列模式不应在 Web 请求内执行 RAG 索引")

    monkeypatch.setattr(JobHuntingApp, "index_rag_long_texts", fail_if_called)
    queue = FakeQueue()
    web_app = create_web_app(resume_dir=tmp_path / "resume-files", task_queue=queue)
    client = TestClient(web_app)
    register_and_login(client, "async-upload@example.com")
    candidate_id = create_profile_over_web(client)

    response = client.post(
        "/api/resumes/upload",
        data={"candidate_id": str(candidate_id)},
        files={"file": ("resume.docx", build_docx_bytes("Python 后端开发经历", "使用 FastAPI 完成求职助手接口项目"), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["indexing_async"] is True
    assert payload["rag_update"] is None
    assert payload["task"]["status"] == "queued"
    assert queue.enqueued == [payload["task"]["task_key"]]
    task_response = client.get(f"/api/tasks/{payload['task']['task_key']}")
    assert task_response.status_code == 200
    assert task_response.json()["task"]["task_type"] == "rag_index"


def test_web_upload_enqueues_ocr_for_scanned_pdf_when_worker_is_enabled(tmp_path) -> None:
    """队列模式下扫描 PDF 上传只创建 OCR 任务，不同步加载 OCR 模型。"""

    class FakeQueue:
        """记录投递任务键的最小 Worker 队列替身。"""

        def __init__(self) -> None:
            self.enqueued: list[str] = []

        def health_check(self) -> None:
            """测试替身始终可用。"""

        def enqueue(self, task_key: str) -> None:
            """记录一次投递。"""

            self.enqueued.append(task_key)

    queue = FakeQueue()
    web_app = create_web_app(resume_dir=tmp_path / "resume-files", task_queue=queue)
    client = TestClient(web_app)
    register_and_login(client, "async-ocr-upload@example.com")
    candidate_id = create_profile_over_web(client)

    response = client.post(
        "/api/resumes/upload",
        data={"candidate_id": str(candidate_id)},
        files={"file": ("scan.pdf", build_scanned_pdf_bytes(), "application/pdf")},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["workflow"] == "ocr"
    assert payload["processing_async"] is True
    assert payload["indexing_async"] is False
    assert payload["artifact"]["status"] == "processing"
    assert payload["artifact"]["long_text_id"] is None
    assert payload["task"]["task_type"] == "resume_ocr"
    assert queue.enqueued == [payload["task"]["task_key"]]


def test_web_tailor_enqueues_resume_export_when_worker_is_enabled(tmp_path) -> None:
    """队列模式下网页只登记简历导出任务，不在 Web 请求中初始化模型。"""

    class FakeQueue:
        """记录任务键的最小 Worker 队列替身。"""

        def __init__(self) -> None:
            self.enqueued: list[str] = []

        def health_check(self) -> None:
            """测试替身始终可用。"""

        def enqueue(self, task_key: str) -> None:
            """记录一次任务投递。"""

            self.enqueued.append(task_key)

    queue = FakeQueue()
    web_app = create_web_app(
        resume_dir=tmp_path / "resume-files",
        task_queue=queue,
        resume_llm_client=StaticLLMClient("Web 不应在请求中调用模型"),
    )
    client = TestClient(web_app)
    register_and_login(client, "async-tailor@example.com")
    candidate_id = create_profile_over_web(client)
    job = client.post(
        "/api/jobs",
        json={
            "raw_text": "Python 后端开发工程师\n职位描述：负责 FastAPI 后端接口开发。",
        },
    ).json()["job"]
    upload_response = client.post(
        "/api/resumes/upload",
        data={"candidate_id": str(candidate_id)},
        files={
            "file": (
                "resume.docx",
                build_docx_bytes(
                    "Python 后端项目经历，负责 FastAPI 接口开发和 PostgreSQL 数据处理。"
                ),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload_response.status_code == 200, upload_response.text
    uploaded = upload_response.json()["artifact"]

    response = client.post(
        f"/api/resumes/{uploaded['id']}/tailor",
        json={"job_id": job["id"], "use_rag": False},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["processing_async"] is True
    assert payload["task"]["task_type"] == "resume_export"
    assert payload["task"]["status"] == "queued"
    assert queue.enqueued[-1] == payload["task"]["task_key"]
    assert len(queue.enqueued) == 2


def test_tailored_resume_creates_docx_and_pdf_without_overwriting_profile_or_source(tmp_path) -> None:
    """职位改写应产出独立草稿和可下载文件，原简历与结构化档案保持不变。"""

    app = JobHuntingApp( resume_dir=tmp_path / "resume-files")
    app.initialize()
    account = app.auth.register("resume@example.com", "password-123")
    candidate_id = app.save_candidate_profile(profile_input(), account_id=account.id)
    job = app.import_job_text(
        """
        Python 后端开发工程师
        15-20K
        杭州
        1-3年
        本科
        职位描述：负责 Python 与 FastAPI 后端接口开发。
        """,
        account_id=account.id,
    )
    original = build_docx_bytes("小林", "Python 与 FastAPI 项目经历")
    source = app.upload_resume_document(
        candidate_id,
        "resume.docx",
        original,
        account_id=account.id,
    )

    generated = app.create_tailored_resume_from_artifact(
        candidate_id=candidate_id,
        source_artifact_id=source.id,
        job_id=job.id,
        llm_client=StaticLLMClient(
            "# 小林\n\n## 求职目标\nPython 后端开发工程师\n\n## 项目经历\n- 使用 Python 与 FastAPI 开发求职助手。"
        ),
        account_id=account.id,
    )

    assert generated.draft.version == 1
    assert {artifact.media_type for artifact in generated.artifacts} == {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/pdf",
    }
    assert all(artifact.parent_artifact_id == source.id for artifact in generated.artifacts)
    assert app.resume_file_path(source).read_bytes() == original
    assert app.get_candidate_profile(candidate_id, account_id=account.id).skills == {"Python": "项目使用"}

    generated_bytes = {
        artifact.media_type: app.resume_file_path(artifact).read_bytes()
        for artifact in generated.artifacts
    }
    assert generated_bytes["application/pdf"].startswith(b"%PDF")
    assert generated_bytes[
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ].startswith(b"PK")
    exported_docx = Document(BytesIO(generated_bytes[
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]))
    assert "Python 后端开发工程师" in "\n".join(paragraph.text for paragraph in exported_docx.paragraphs)

    deleted_source = app.delete_resume_artifact(source.id, account_id=account.id)
    remaining_artifacts = app.list_resume_artifacts(candidate_id, account_id=account.id)

    assert deleted_source["artifact_id"] == source.id
    assert [artifact.id for artifact in remaining_artifacts] == [item.id for item in generated.artifacts]
    assert app.resume_file_path(generated.artifacts[0]).read_bytes() == generated_bytes[generated.artifacts[0].media_type]


def test_web_can_tailor_uploaded_resume_and_return_download_urls(tmp_path, monkeypatch) -> None:
    """网页职位定制接口应返回独立草稿和两个可直接下载的文件版本。"""

    def fail_search(*_args, **_kwargs):
        raise AssertionError("Web 禁用 RAG 时不应执行语义检索")

    monkeypatch.setattr(JobHuntingApp, "search_rag", fail_search)

    web_app = create_web_app(
        resume_dir=tmp_path / "resume-files",
        resume_llm_client=StaticLLMClient(
            "# 小林\n\n## 求职目标\nPython 后端开发工程师\n\n## 项目经历\n- 使用 Python 与 FastAPI 开发求职助手。"
        ),
    )
    client = TestClient(web_app)
    register_and_login(client, "tailor@example.com")
    candidate_id = create_profile_over_web(client)
    job = client.post(
        "/api/jobs",
        json={
            "raw_text": """
            Python 后端开发工程师
            15-20K
            杭州
            1-3年
            本科
            职位描述：负责 Python 与 FastAPI 后端接口开发。
            """,
        },
    ).json()["job"]
    uploaded = client.post(
        "/api/resumes/upload",
        data={"candidate_id": str(candidate_id)},
        files={
            "file": (
                "resume.docx",
                build_docx_bytes("小林", "Python 与 FastAPI 项目经历"),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    ).json()["artifact"]

    tailored = client.post(
        f"/api/resumes/{uploaded['id']}/tailor",
        json={"job_id": job["id"], "use_rag": False},
    )

    assert tailored.status_code == 200
    payload = tailored.json()
    assert payload["draft"]["status"] == "需候选人确认"
    assert len(payload["artifacts"]) == 2
    assert all(item["parent_artifact_id"] == uploaded["id"] for item in payload["artifacts"])
    for artifact in payload["artifacts"]:
        downloaded = client.get(artifact["download_url"])
        assert downloaded.status_code == 200
        assert downloaded.content

    generated_ids = [item["id"] for item in payload["artifacts"]]
    deleted_tailored = client.delete(f"/api/resumes/{generated_ids[0]}")
    remaining = client.get("/api/resumes", params={"candidate_id": candidate_id})

    assert deleted_tailored.status_code == 200
    assert deleted_tailored.json()["deleted"] is True
    assert [item["id"] for item in remaining.json()["artifacts"]] == [uploaded["id"], generated_ids[1]]


def test_vue_frontend_exposes_resume_upload_tailor_and_download_workflow(tmp_path) -> None:
    """Vue 工作台应提供完整文件交互，而不是只存在不可发现的后端 API。"""

    client = TestClient(
        create_web_app(
            resume_dir=tmp_path / "resume-files",
        )
    )

    home = client.get("/").text
    script = client.get("/static/app.js").text
    styles = client.get("/static/styles.css").text

    assert 'ref="resumeFileInput"' in home
    assert 'accept=".docx,.pdf"' in home
    assert '@change="uploadResume"' in home
    assert "简历文件" in home
    assert "生成定制版" in home
    assert ':href="artifact.download_url"' in home
    assert "async uploadResume(event)" in script
    assert '"/api/resumes/upload"' in script
    assert 'task.task_type === "resume_ocr"' in script
    assert "rag_task_key" in script
    assert "async loadResumeArtifacts(signal = null)" in script
    assert "async tailorResume(artifact)" in script
    assert "async deleteResumeArtifact(artifact)" in script
    assert "/api/resumes/${encodeURIComponent(artifact.id)}" in script
    assert "resume-delete" in home
    assert "artifact.status !== 'ready'" in home
    assert "resume-artifact-list" in styles
